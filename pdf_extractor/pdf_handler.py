from PyPDF2 import PdfReader
from docx import Document
import textract
import win32com.client
import subprocess
import pythoncom
import os

def extract_text_from_pdf(file_path):
    pdf_document = PdfReader(file_path)
    text = ""
    for page in pdf_document.pages:
        text += page.extract_text()
    print(text)    
    return text

def extract_text_from_docx(file_path):
    
    doc = Document(file_path)
    text = ""

    try:
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text.strip() + "\n"
            

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        # Extract header
        header = doc.sections[0].header
        for para in header.paragraphs:
            text += "Header: " + para.text + "\n"

        # Extract footer
        footer = doc.sections[0].footer
        for para in footer.paragraphs:
            text += "Footer: " + para.text + "\n"

        print(text)
        return text
    
    except Exception as e:
        print(f"Error extracting text from .docx file: {e}")
        return None


def extract_text_from_doc(file_path):
    try:
        text = textract.process(file_path)
        print(text)
        return text.decode('utf-8')
    except Exception as e:
        print(f"Error extracting text from .doc file: {e}")
        return None


# def extract_text_from_doc(file_path):
#     if not os.path.exists(file_path):
#         raise FileNotFoundError(f"File not found: {file_path}")

#     pythoncom.CoInitialize()
    
#     try:
#         word = win32com.client.Dispatch("Word.Application")
#         word.Visible = False

#         try:
#             doc = word.Documents.Open(file_path)
#             text = doc.Content.Text
#             print(text)
#             doc.Close()
#             return text
#         except Exception as e:
#             raise Exception(f"Failed to open document: {e}")
#     finally:
#         word.Quit()
#         pythoncom.CoUninitialize()


# def extract_text_from_doc(file_path):
#     # Ensure the file exists
#     if not os.path.exists(file_path):
#         raise FileNotFoundError(f"File not found: {file_path}")

#     # Initialize COM library
#     pythoncom.CoInitialize()

#     try:
#         word = win32com.client.Dispatch("Word.Application")
#         word.Visible = False
        
#         # Try opening the document
#         try:
#             doc = word.Documents.Open(file_path)
#         except Exception as e:
#             raise Exception(f"Failed to open document: {e}")

#         # Extract text
#         try:
#             text = doc.Content.Text
#         except Exception as e:
#             raise Exception(f"Failed to extract text: {e}")
#         finally:
#             doc.Close()

#         word.Quit()
#         return text

#     except Exception as e:
#         print(f"Error: {e}")
#     finally:
#         # Uninitialize COM library
#         pythoncom.CoUninitialize()
